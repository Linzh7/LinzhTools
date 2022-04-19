import re
from docx import Document

path = 'xxx.docx'
newPath = 'new_' + path
citePattern = r"\[[-,\d]+\]"
citesReplaceDic = {}


def replaceByDictionary(matched):
    if matched:
        string = matched.group()[1:-1]
        citeIndexList = citesStringDivide(string)
        for i in range(len(citeIndexList)):
            citeIndexList[i] = citesReplaceDic.get(int(citeIndexList[i]))
        return re.sub(r"[' ]", "", str(citeIndexList))


def citesStringDivide(citesString):
    cites = []
    if ',' in citesString:
        citesElementList = citesString.split(',')
    else:
        citesElementList = [citesString]
    for citesElement in citesElementList:
        if '-' in citesElement:
            citesElementSplitList = citesElement.split('-')
            cites += list(range(int(citesElementSplitList[0]), int(citesElementSplitList[1])+1))
        else:
            cites += [citesElement]
    return cites


if __name__ == '__main__':
    citesIndexDic = {}
    citesDocIndex = 1
    Document = Document(path)

    for paragraph in Document.paragraphs:
        citesStringList = re.findall(citePattern, paragraph.text)
        for citesString in citesStringList:
            citesString = citesString[1:-1]
            cites = citesStringDivide(citesString)
            for cite in cites:
                if int(cite) in citesIndexDic.keys():
                    continue
                citesIndexDic[int(cite)] = citesDocIndex
                citesDocIndex += 1
    # print(citesIndexDic)
    MAXN = citesDocIndex * 5

    for key, value in citesIndexDic.items():
        citesReplaceDic[key] = value + MAXN

    for paragraph in Document.paragraphs:
        paragraph.text = re.sub(citePattern, replaceByDictionary, paragraph.text)

    tmpDic = {}
    for key, value in citesReplaceDic.items():
        tmpDic[value] = value - MAXN
    citesReplaceDic = tmpDic

    for paragraph in Document.paragraphs:
        paragraph.text = re.sub(citePattern, replaceByDictionary, paragraph.text)

    Document.save(newPath)
